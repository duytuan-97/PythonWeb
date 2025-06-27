# Fixx delete photo: kiểm tra hình ảnh, không xóa hình ảnh của minh chứng dùng chung khi xóa minh chứng
from datetime import datetime
import os
import shutil
from django.contrib import messages as dj_messages
from django import forms
from django.contrib import admin
from django.shortcuts import redirect, render
from django.db.models import Count

from CTDT.forms import AttestForm, CommonAttestForm
from CTDT.image_utils import remove_image_from_index
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.section import WD_SECTION
from functools import lru_cache
# from CTDT.model_train.ml_model import predict_image, train_model


# from CTDT import forms

# Register your models here.
from .models import PhotoCommonAttest, Post, PhotoAttest
from .models import box
from .models import standard
from .models import criterion
from .models import attest
from .models import common_attest

from django_admin_listfilter_dropdown.filters import DropdownFilter, RelatedDropdownFilter, ChoiceDropdownFilter

from import_export.admin import ImportExportActionModelAdmin

from django.utils.html import format_html
from django.urls import reverse
from django.utils.http import urlencode

from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

from collections import defaultdict

from django.utils.html import format_html

from django.urls import path

from django.conf import settings
from django.contrib.auth import get_user_model
from .notifications import EmailNotification
from .admin_convert.action_convert import ActionConvert
from django.contrib.admin.widgets import AdminTextInputWidget 

from django.utils.text import slugify
from django.template.loader import get_template
from django.utils.translation import gettext as _

from easy_thumbnails.files import get_thumbnailer
from django.db import transaction


from django.contrib.auth.admin import UserAdmin as BaseUserAdmin
from .models import ProfileUser

from .image_utils import search_similar_images
from django.core.exceptions import ValidationError

from guardian.admin import GuardedModelAdmin
from guardian.shortcuts import get_objects_for_user, remove_perm


# from django.utils.decorators import method_decorator
# from django.contrib.admin.views.decorators import staff_member_required

# admin.site.register(Post)

User = get_user_model()

class SendMailInline(admin.StackedInline):
    model = ProfileUser
    sendMailUser = False
    

# Define a new User admin
class UserAdmin(BaseUserAdmin):
    inlines = (SendMailInline, )

# Re-register UserAdmin
admin.site.unregister(User)
admin.site.register(User, UserAdmin)


@admin.action(description="Mark selected stories as published")
def make_published(modeladmin, request, queryset):
    queryset.update(status="p")

# hộp
# admin.site.register(box)
@admin.register(box)
class boxAdmin(admin.ModelAdmin): 
    search_fields = ('title',)
    # prepopulated_fields = {'slug': ['title']}

    class Media:
        js = (['../static/js/custom_admin.js', 'https://cdnjs.cloudflare.com/ajax/libs/speakingurl/14.0.1/speakingurl.min.js'])  # Đường dẫn file JS

    def save_model(self, request, obj, form, change):
        super().save_model(request, obj, form, change)  # 🔹 Đảm bảo obj đã được lưu trước khi lấy pk
        if change :
            action_type = "Cập nhật hộp"
        else :
            action_type = "Thêm mới hộp"
        admin_url = request.build_absolute_uri(reverse('admin:CTDT_box_change', args=[obj.pk]))
        
        # EmailNotification.send_box_email(request, [obj], action_type, admin_url)

    def delete_model(self, request, obj):
        # EmailNotification.send_box_email(request, [obj], "Xóa hộp", "Delete")
        super().delete_model(request, obj)
    
    def delete_queryset(self, request, queryset):
        # EmailNotification.send_box_email(request, queryset, "Xóa hộp", "Delete")
        super().delete_queryset(request, queryset)
        
# tiêu chuẩn
# admin.site.register(standard)
@admin.register(standard)
class standardAdmin(admin.ModelAdmin):
    list_display = ('title','view_criterion_link',) 
    search_fields = ('title',)
    ordering = ('title',)
    # prepopulated_fields = {'slug': ['title']}
    
    class Media:
        js = (['../static/js/custom_admin.js', 'https://cdnjs.cloudflare.com/ajax/libs/speakingurl/14.0.1/speakingurl.min.js'])  # Đường dẫn file JS
        
    
    def view_criterion_link(self, obj):
        count = obj.criterion_set.count()
        url = (
            reverse("admin:CTDT_criterion_changelist")#tên Project_model_linkChange
            + "?"
            + urlencode({"standard__id": f"{obj.id}"})
        )
        return format_html('<a href="{}">{} Tiêu chí</a>', url, count)
    
    def save_model(self, request, obj, form, change):
        super().save_model(request, obj, form, change)  # 🔹 Đảm bảo obj đã được lưu trước khi lấy pk
        if change :
            action_type = "Cập nhật tiêu chuẩn"
            obj.slug = slugify(obj.title)  # Cập nhật lại slug từ title
        else :
            action_type = "Thêm mới tiêu chuẩn"
        admin_url = request.build_absolute_uri(reverse('admin:CTDT_standard_change', args=[obj.pk]))
        super().save_model(request, obj, form, change)
        # EmailNotification.send_standard_email(request, [obj], action_type, admin_url)

    def delete_model(self, request, obj):
        # EmailNotification.send_standard_email(request, [obj], "Xóa tiêu chuẩn", "Delete")
        super().delete_model(request, obj)
    
    def delete_queryset(self, request, queryset):
        # EmailNotification.send_standard_email(request, queryset, "Xóa tiêu chuẩn", "Delete")
        super().delete_queryset(request, queryset)

    view_criterion_link.short_description = "Tiêu chí"
    
    
# tiêu chí
#admin.site.register(criterion)
@admin.register(criterion)
# class criterionAdmin(admin.ModelAdmin):   
class criterionAdmin(GuardedModelAdmin):   
    list_display = ('standard_name','title', 'view_attests_link',)
    list_display_links = ('title',)
    ordering = ('standard','title',)
    list_filter = (
        #'standard',
        # for ordinary fields
        #('standard', DropdownFilter),
        # for choice fields
        #('standard', ChoiceDropdownFilter),
        # for related fields
        ('standard', RelatedDropdownFilter),
    )
    
    class Media:
        js = (['https://code.jquery.com/jquery-3.6.0.min.js','../static/js/custom_admin.js', 'https://cdnjs.cloudflare.com/ajax/libs/speakingurl/14.0.1/speakingurl.min.js'])  # Đường dẫn file JS
        
    search_fields = ('title',)
    prepopulated_fields = {'slug': ['title']}
    
    @admin.display(description="Tiêu chuẩn")
    def standard_name(self, obj):
        if self.model.objects.filter(standard_id=obj.standard, pk__lt=obj.pk).exists():
            return ""
        return f"{obj.standard}".upper()
    
    def view_attests_link(self, obj):
        count = obj.attest_set.count()
        url = (
            reverse("admin:CTDT_attest_changelist")#tên Project_model_linkChange
            + "?"
            + urlencode({"criterion__id": f"{obj.id}"})
        )
        return format_html('<a href="{}">{}</a>', url, count)
    
    def save_model(self, request, obj, form, change):
        super().save_model(request, obj, form, change)  # 🔹 Đảm bảo obj đã được lưu trước khi lấy pk
        if change :
            action_type = "Cập nhật tiêu chí"
        else :
            action_type = "Thêm mới tiêu chí"
        admin_url = request.build_absolute_uri(reverse('admin:CTDT_criterion_change', args=[obj.pk]))
        # EmailNotification.send_criterion_email(request, [obj], action_type, admin_url)

    def delete_model(self, request, obj):
        # EmailNotification.send_criterion_email(request, [obj], "Xóa tiêu chí", "Delete")
        super().delete_model(request, obj)
    
    def delete_queryset(self, request, queryset):
        # EmailNotification.send_criterion_email(request, queryset, "Xóa tiêu chí", "Delete")
        super().delete_queryset(request, queryset)

    view_attests_link.short_description = "Minh chứng"
    
    # ===================phân quyền===============================
    def has_module_permission(self, request):
        if super().has_module_permission(request):
            return True
        return self.get_model_objects(request).exists()
    
    # def get_queryset(self, request):
    #     if request.user.is_superuser:
    #         return super().get_queryset(request)
    #     data = self.get_model_objects(request)
    #     return data
    
    def get_queryset(self, request):
        print(f"Getting queryset for user: {request.user}, is_superuser: {request.user.is_superuser}")
        if request.user.is_superuser:
            qs = super().get_queryset(request)
            print(f"Superuser queryset: {qs.count()} objects")
            return qs
        data = self.get_model_objects(request)
        print(f"Filtered queryset: {data.count()} objects")
        return data
    
    def get_model_objects(self, request, action=None, klass=None):
        opts = self.opts
        actions = [action] if action else ['view', 'change', 'delete']
        klass = klass if klass else opts.model
        model_name = klass._meta.model_name
        
        return get_objects_for_user(user=request.user, perms=[f'{perm}_{model_name}' for perm in actions], 
                                    klass=klass, any_perm=True, accept_global_perms=False)
    
    def has_permission(self, request, obj, action):
        opts = self.opts
        code_name = f'{action}_{opts.model_name}'
        if obj:
            return request.user.has_perm(f'{opts.app_label}.{code_name}', obj)
        else:
            return True
            # return self.get_model_objects(request).exists()
        
    def has_view_permission(self, request, obj = None):
        return self.has_permission(request, obj, 'view')
        # return True
    
    def has_change_permission(self, request, obj = None):
        # result = self.has_permission(request, obj, 'change')
        return self.has_permission(request, obj, 'change')
        # print(request.user.is_authenticated)
        # print(get_objects_for_user(request.user, 'CTDT.change_criterion'))
        # print(f"User: {request.user}, Obj: {obj}, Permission: {result}")
        # return result
        # return True
    
    def has_delete_permission(self, request, obj = None):
        return self.has_permission(request, obj, 'delete')
        # return True
# admin.site.register(criterion, criterionAdmin)
class PhotoAttestInline(admin.TabularInline):
    model = PhotoAttest
    fields = ("showphoto_thumbnail",)
    readonly_fields = ("showphoto_thumbnail",)
    max_num = 0

    def showphoto_thumbnail(self, instance):
        """A (pseudo)field that returns an image thumbnail for a show photo."""
        tpl = get_template("admin/templates/show_thumbnail.html")
        return tpl.render({"photo": instance.photo})
    
    def clean_photo(self, instance):
        if instance.photo:
            # instance.clean()  # Gọi phương thức clean của model để kiểm tra ảnh trùng lặp
            similar_images = search_similar_images(self.photo.path)
            if similar_images:
                raise ValidationError(f"Hình ảnh này có thể đã tồn tại: {', '.join([img[0] for img in similar_images])}")

    showphoto_thumbnail.short_description = _("Thumbnail")

class PhotoCommonAttestInline(admin.TabularInline):
    model = PhotoCommonAttest
    fields = ("showphoto_thumbnail",)
    readonly_fields = ("showphoto_thumbnail",)
    max_num = 0

    def showphoto_thumbnail(self, instance):
        """A (pseudo)field that returns an image thumbnail for a show photo."""
        tpl = get_template("admin/templates/show_thumbnail.html")
        return tpl.render({"photo": instance.photo})

    def clean_photo(self, instance):
        if instance.photo:
            # instance.clean()  # Gọi phương thức clean của model để kiểm tra ảnh trùng lặp
            similar_images = search_similar_images(self.photo.path)
            if similar_images:
                raise ValidationError(f"Hình ảnh này có thể đã tồn tại: {', '.join([img[0] for img in similar_images])}")
    
    showphoto_thumbnail.short_description = _("Thumbnail")

#admin.site.register(attest)
@admin.register(attest)
class attestAdmin(admin.ModelAdmin):
# class attestAdmin(ImportExportActionModelAdmin, admin.ModelAdmin):
    
    change_list_template = "admin/CTDT/attest/change_list.html"
    
    form = AttestForm
    inlines = [PhotoAttestInline]

    def save_related(self, request, form, formsets, change):
        super().save_related(request, form, formsets, change)
        form.save_photos(form.instance)
    # exclude = {'is_common',}
    list_display = ('criterion_name', 'attest_id_name','attest_stt', 'title', 'body', 'performer', "is_common")
    list_display_links = ('title',)
    list_filter = (
        #'criterion',
        # for ordinary fields
        #('criterion', DropdownFilter),
        # for choice fields
        #('criterion', ChoiceDropdownFilter),
        # for related fields
        ('criterion', RelatedDropdownFilter),
    )
    
    ordering = ('criterion_id', 'attest_id', 'attest_stt')
    search_fields = ('title', 'performer')
    # prepopulated_fields = {'slug': ['attest_id','attest_stt']}
    
    # def get_readonly_fields(self, request, obj=None):
    #     if obj:  # Trang chỉnh sửa (change)
    #         return self.readonly_fields + ('attest_id',)
    #     return self.readonly_fields
    
    def get_form(self, request, obj=None, **kwargs):
        form = super().get_form(request, obj, **kwargs)
        if obj:
            if obj.common_attest is not None:  # Nếu là minh chứng dùng chung
                for field_name in form.base_fields:
                    form.base_fields[field_name].disabled = True  # Vô hiệu hóa trường
            else:
                form.base_fields['common_attest'].disabled = True
                form.base_fields['is_common'].disabled = True
                if 'photos' in form.base_fields:
                    form.base_fields['photos'].disabled = False
        class FormWithRequest(form):
            def __init__(self2, *args, **inner_kwargs):
                inner_kwargs['request'] = request
                super().__init__(*args, **inner_kwargs)
        return FormWithRequest
    
    def render_change_form(self, request, context, add=False, change=False, form_url="", obj=None):
        if obj and obj.common_attest is not None:  # Nếu là minh chứng dùng chung
            context['show_save'] = False  # Ẩn nút lưu
            context['show_save_and_continue'] = False
            context['show_save_and_add_another'] = False
        else:
            context['show_save'] = True
            context['show_save_and_continue'] = True
            context['show_save_and_add_another'] = True
        return super().render_change_form(request, context, add, change, form_url, obj)
    

    def save_model(self, request, obj, form, change):
        
        super().save_model(request, obj, form, change)  # 🔹 Đảm bảo obj đã được lưu trước khi lấy pk
        if change :
            action_type = "Cập nhật minh chứng"
        else :
            action_type = "Thêm mới minh chứng"
        admin_url = request.build_absolute_uri(reverse('admin:CTDT_attest_change', args=[obj.pk]))
        
        
        if not change and obj.common_attest : 
            # obj.is_common = bool(obj.common_attest)  # Gán True nếu common_attest có giá trị
            common_attest_data = obj.common_attest
            obj.attest_id = common_attest_data.common_attest_id
            obj.attest_stt = common_attest_data.common_attest_stt
            obj.title = common_attest_data.title
            obj.body = common_attest_data.body
            obj.performer = common_attest_data.performer
            obj.note = "DC"
            obj.slug = common_attest_data.slug
            # obj.image = common_attest_data.image
            # obj.criterion = common_attest_data.criterion
            obj.box = common_attest_data.box
            obj.is_common = True
            
        else:
            obj.is_common = False
        
        
        user = request.user
        send_mail_obj, created = ProfileUser.objects.get_or_create(user=user)
        if send_mail_obj.SendMailUser:
            dj_messages.success(request, f"✅ Check send mail user {send_mail_obj.SendMailUser}")
        # # EmailNotification.send_attest_email(request, [obj], action_type, admin_url)
        # EmailNotification.send_attest_email(request, [obj], action_type)
        transaction.on_commit(lambda: 
            EmailNotification.send_attest_email(request, [obj], action_type)
        )
        super().save_model(request, obj, form, change)
        
        # Xử lý các PhotoAttest instances
        if 'photoattest_set' in form.cleaned_data:  # Kiểm tra xem có PhotoAttest instances trong form không
            for photo in form.cleaned_data['photoattest_set']:
                try:
                    self.inlines[0].clean_photo(photo)  # Gọi hàm clean_photo nếu cần
                except forms.ValidationError as e:
                    raise forms.ValidationError(e)  # Ném lỗi nếu có vấn đề
                photo.save()  # Lưu PhotoAttest instance
    
    def delete_model(self, request, obj):
        """ Gửi email khi xóa """
        
        user = request.user
        send_mail_obj, created = ProfileUser.objects.get_or_create(user=user)
        if send_mail_obj.SendMailUser:
            dj_messages.success(request, f"✅ Check send mail user {send_mail_obj.SendMailUser}")
        
        # # EmailNotification.send_attest_email(request, [obj], "Xóa minh chứng", "Delete")
        EmailNotification.send_attest_email(request, [obj], "Xóa minh chứng")
        
        super().delete_model(request, obj)
    
    def delete_queryset(self, request, queryset):
        """ Gửi email chứa danh sách attest bị xóa trước khi xóa """
        # delete photo
        for attest in queryset:
            for photo_attest in attest.photos.all():  # Lấy ảnh liên kết
                if photo_attest.photo and not common_attest:
                    
                    # # Kiểm tra xem ảnh có được sử dụng bởi các attest khác hay không
                    # photo_usage_count = photo_attest.photo.attestphoto_set.aggregate(Count('id'))['id__count']
                    
                    # # Chỉ xóa ảnh nếu ảnh không được sử dụng bởi attest nào khác
                    # if photo_usage_count <= 1:
                    try:
                        thumbnail_path = get_thumbnailer(photo_attest.photo)['small'].path
                        if os.path.isfile(thumbnail_path):
                            os.remove(thumbnail_path)
                    except Exception:
                        pass  # Bỏ qua nếu không có thumbnail

                    if os.path.isfile(photo_attest.photo.path):
                        os.remove(photo_attest.photo.path)
                    folder = os.path.dirname(photo_attest.photo.path)
                    # xóa index
                    remove_image_from_index(photo_attest.photo.path)
                    # Kiểm tra và xóa folder nếu rỗng (loại bỏ file ẩn nếu cần)
                    remaining_files = [f for f in os.listdir(folder) if not f.startswith('.')]
                    if not remaining_files:
                        shutil.rmtree(folder)
        
        user = request.user
        send_mail_obj, created = ProfileUser.objects.get_or_create(user=user)
        if send_mail_obj.SendMailUser:
            dj_messages.success(request, f"✅ Check send mail user {send_mail_obj.SendMailUser}")
        
        # EmailNotification.send_attest_email(request, queryset, "Xóa minh chứng", "Delete")
        EmailNotification.send_attest_email(request, queryset, "Xóa minh chứng")
        
        # Gọi phương thức mặc định để xóa các attest
        super().delete_queryset(request, queryset)
    
    class Media:
        js = ('../static/js/custom_admin.js', 'https://cdnjs.cloudflare.com/ajax/libs/speakingurl/14.0.1/speakingurl.min.js')  # Đường dẫn file JS
        css = {
            'all': ('../static/css/custom_admin.css',)
        }
        

    # @admin.display(description="Tiêu chí")
    # def criterion_name(self, obj):
    #     if self.model.objects.filter(criterion=obj.criterion, pk__lt=obj.pk).exists():
    #         return ""
    #     # return f"{obj.criterion} {obj.title}".upper()
    #     return f"{obj.criterion}"
    # @admin.display(description="Mã minh chứng")
    # def attest_id_name(self, obj):
    #     if self.model.objects.filter(attest_id=obj.attest_id, pk__lt=obj.pk).exists():
    #         return ""
    #     return f"{obj.attest_id}".upper()
    
    def get_queryset(self, request):
        qs = super().get_queryset(request)
        self._displayed_criteria_attest = set()  # Store seen criterion_ids
        return qs.order_by('criterion__standard', 'criterion', 'attest_id')

    @admin.display(description="Tiêu chí")
    def criterion_name(self, obj):
        key = obj.criterion_id
        if key in self._displayed_criteria_attest:
            return ""
        self._displayed_criteria_attest.add(key)
        return str(obj.criterion)

    @admin.display(description="Mã minh chứng")
    def attest_id_name(self, obj):
        if self.model.objects.filter(
            attest_id=obj.attest_id,
            criterion=obj.criterion,
            pk__lt=obj.pk
        ).exists():
            return ""
        return obj.attest_id.upper()
    
    
    
    #  # Thêm đường dẫn cho action nhập file Word
    # def get_urls(self):
    #     urls = super().get_urls()
    #     custom_urls = [
    #         path('import-word/', self.import_word, name='import_word'),
    #     ]
    #     return custom_urls + urls
    
    # # Giao diện tải file Word
    # def import_word(self, request):
    #     from django.template.response import TemplateResponse
    #     from django.core.files.uploadedfile import UploadedFile

    #     if request.method == "POST" and request.FILES.get('word_file'):
    #         word_file = request.FILES['word_file']
    #         if isinstance(word_file, UploadedFile):
    #             self.handle_uploaded_word(word_file)
    #             self.message_user(request, "File Word đã được xử lý.")
    #             return HttpResponse(".")
    #     return TemplateResponse(request, "admin/import_word.html", context={})
    
    
    actions = ['export_to_word']
    # def export_to_word(self, request, queryset):

    #     # Tạo file Word mới
    #     document = Document()

    #     # Thêm tiêu đề chính
    #     title = document.add_paragraph("DANH MỤC MINH CHỨNG")
    #     title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #     title.runs[0].font.size = Pt(14)
    #     title.runs[0].bold = True

    #     # Tạo bảng với 5 cột
    #     table = document.add_table(rows=1, cols=5)
    #     table.style = 'Table Grid'

    #     # Định nghĩa header
    #     hdr_cells = table.rows[0].cells
    #     hdr_cells[0].text = 'Tiêu chí'
    #     hdr_cells[1].text = 'STT (trong tiêu chí)'
    #     hdr_cells[2].text = 'Mã minh chứng'
    #     hdr_cells[3].text = 'Tên minh chứng'
    #     hdr_cells[4].text = 'Nơi ban hành...'

    #     # Biến theo dõi tiêu chuẩn và tiêu chí hiện tại
    #     current_standard = None
    #     current_criterion = None
    #     current_attest_id = None
    #     merge_criterion_cell = None
    #     merge_attest_cell = None
    #     merge_stt_cell = None
    #     stt_within_criterion = 0

    #     # Sắp xếp dữ liệu theo tiêu chuẩn và tiêu chí
    #     sorted_queryset = queryset.order_by('criterion__standard', 'criterion', 'attest_id')

    #     for obj in sorted_queryset:
    #         # Kiểm tra tiêu chuẩn (standard)
    #         if obj.criterion.standard != current_standard:
    #             # Nếu tiêu chuẩn thay đổi, thêm dòng tiêu chuẩn mới
    #             current_standard = obj.criterion.standard
    #             row_cells = table.add_row().cells
    #             row_cells[0].merge(row_cells[4])  # Gộp tất cả các cột
    #             row_cells[0].text = f"Tiêu chuẩn: {current_standard.title}"
    #             row_cells[0].paragraphs[0].runs[0].bold = True

    #         # Kiểm tra tiêu chí (criterion)
    #         if obj.criterion != current_criterion:
    #             # Nếu tiêu chí thay đổi, gộp ô tiêu chí trước đó (nếu có)
    #             if merge_criterion_cell is not None:
    #                 merge_criterion_cell.merge(row_cells[0])
    #             current_criterion = obj.criterion
    #             merge_criterion_cell = None
    #             stt_within_criterion = 0  # Đặt lại bộ đếm STT

    #         # Kiểm tra mã minh chứng (attest_id)
    #         if obj.attest_id != current_attest_id:
    #             # Nếu mã minh chứng thay đổi, gộp ô minh chứng trước đó (nếu có)
    #             if merge_attest_cell is not None:
    #                 merge_attest_cell.merge(row_cells[2])
    #             if merge_stt_cell is not None:
    #                 merge_stt_cell.merge(row_cells[1])
    #             current_attest_id = obj.attest_id
    #             merge_attest_cell = None
    #             merge_stt_cell = None
    #             stt_within_criterion += 1

    #         # Thêm dòng chi tiết
    #         row_cells = table.add_row().cells

    #         # Gán giá trị cho cột "Tiêu chí"
    #         if current_criterion == obj.criterion and merge_criterion_cell is None:
    #             merge_criterion_cell = row_cells[0]
    #             row_cells[0].text = str(obj.criterion)
    #         else:
    #             row_cells[0].text = ''

    #         # Gán giá trị cho cột "STT trong tiêu chí"
    #         if current_attest_id == obj.attest_id and merge_stt_cell is None:
    #             merge_stt_cell = row_cells[1]
    #             row_cells[1].text = str(stt_within_criterion)
    #         else:
    #             row_cells[1].text = ''

    #         # Gán giá trị cho cột "Mã minh chứng"
    #         if current_attest_id == obj.attest_id and merge_attest_cell is None:
    #             merge_attest_cell = row_cells[2]
    #             row_cells[2].text = obj.attest_id
    #         else:
    #             row_cells[2].text = ''

    #         # Gán giá trị cho các cột khác
    #         row_cells[3].text = obj.title  # Tên minh chứng
    #         row_cells[4].text = obj.performer  # Nơi ban hành

    #     # Gộp ô cuối cùng nếu còn dư
    #     if merge_criterion_cell is not None:
    #         merge_criterion_cell.merge(row_cells[0])
    #     if merge_attest_cell is not None:
    #         merge_attest_cell.merge(row_cells[2])
    #     if merge_stt_cell is not None:
    #         merge_stt_cell.merge(row_cells[1])

    #     # Lưu tài liệu vào HTTP response
    #     response = HttpResponse(
    #         content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    #     )
    #     response['Content-Disposition'] = 'attachment; filename="DanhMucMinhChung.docx"'
    #     document.save(response)
    #     return response
    
    def export_to_word(self, request, queryset):
        document = Document()
        
        # === Đóng khung section 0 ===
        section0 = document.sections[0]
        
        section0.left_margin = Inches(0.7)
        section0.right_margin = Inches(0.7)
        pgBorders = OxmlElement('w:pgBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')
            border.set(qn('w:space'), '4')
            border.set(qn('w:color'), 'auto')
            pgBorders.append(border)
        section0._sectPr.append(pgBorders)
        
        # import word
        today = datetime.today()
        current_month = today.strftime("%m")
        current_year = today.strftime("%Y")
        
        document.add_paragraph("\n")
        p = document.add_paragraph("BỘ TÀI NGUYÊN VÀ MÔI TRƯỜNG")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(16)

        p = document.add_paragraph("TRƯỜNG ĐẠI HỌC TÀI NGUYÊN VÀ MÔI TRƯỜNG TP. HCM")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(15)
        run.bold = True
        
        #kẻ đường với bảng
        table = document.add_table(rows=1, cols=1)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = table.cell(0, 0)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell.width = Inches(5.0)  # Độ dài 

        # Thêm đường kẻ dưới ô bảng
        cell_border = OxmlElement('w:pBdr')
        bottom_border = OxmlElement('w:top')
        bottom_border.set(qn('w:val'), 'single')
        bottom_border.set(qn('w:sz'), '12')
        bottom_border.set(qn('w:space'), '6')
        bottom_border.set(qn('w:color'), 'auto')
        cell_border.append(bottom_border)
        cell.paragraphs[0]._element.get_or_add_pPr().append(cell_border)

        # ===== CHÈN LOGO =====
        document.add_paragraph("\n")
        document.add_picture('media/' + 'logo.png', width=Inches(2.0))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ===== TIÊU ĐỀ =====
        p = document.add_paragraph("\nDANH MỤC MINH CHỨNG TỰ ĐÁNH GIÁ")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(20)

        # ===== CHƯƠNG TRÌNH =====
        p = document.add_paragraph("CHƯƠNG TRÌNH ĐÀO TẠO TRÌNH ĐỘ")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(15.5)

        p = document.add_paragraph("ĐẠI HỌC NGÀNH CÔNG NGHỆ THÔNG TIN")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(15.5)

        # ===== TIÊU CHUẨN =====
        p = document.add_paragraph("\nTheo tiêu chuẩn đánh giá chất lượng chương trình đào tạo")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        p.runs[0].font.size = Pt(15.5)

        p = document.add_paragraph("của Bộ Giáo dục và Đào tạo")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        p.runs[0].font.size = Pt(15.5)

        for _ in range(3):
            document.add_paragraph()
        # ===== THỜI GIAN ĐỊA ĐIỂM =====
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run(f"TP Hồ Chí Minh, tháng {int(current_month)} năm {current_year}")
        run.italic = True
        run.font.size = Pt(14)

        document.add_page_break()
        # end import word
        
        # Chèn section mới (Section 1) bắt đầu từ trang tiếp theo
        document.add_section(WD_SECTION.NEW_PAGE)
        section1 = document.sections[1]
        pgBorders1 = OxmlElement('w:pgBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')  # Tắt viền bằng 'none'
            pgBorders1.append(border)
        section1._sectPr.append(pgBorders1)
        # Ngắt liên kết header/footer với section trước
        section1.header.is_linked_to_previous = False
        section1.footer.is_linked_to_previous = False
        # ===== TRANG BÌA LÓT =====
        document.add_paragraph("\n")
        p = document.add_paragraph("BỘ TÀI NGUYÊN VÀ MÔI TRƯỜNG")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(16)

        p = document.add_paragraph("TRƯỜNG ĐẠI HỌC TÀI NGUYÊN VÀ MÔI TRƯỜNG TP. HCM")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(15)
        run.bold = True
        rPr = run._element.get_or_add_rPr()
        
        #kẻ đường với bảng
        table = document.add_table(rows=1, cols=1)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = table.cell(0, 0)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell.width = Inches(5.0)  # Độ dài 

        # Thêm đường kẻ dưới ô bảng
        cell_border = OxmlElement('w:pBdr')
        bottom_border = OxmlElement('w:top')
        bottom_border.set(qn('w:val'), 'single')
        bottom_border.set(qn('w:sz'), '12')
        bottom_border.set(qn('w:space'), '6')
        bottom_border.set(qn('w:color'), 'auto')
        cell_border.append(bottom_border)
        cell.paragraphs[0]._element.get_or_add_pPr().append(cell_border)

        # Logo
        document.add_paragraph("\n")
        document.add_picture('media/logo.png', width=Inches(2.0))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Tiêu đề chính
        p = document.add_paragraph("\nDANH MỤC MINH CHỨNG")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(20)

        # Ghi chú dưới tiêu đề
        p = document.add_paragraph("(Kèm theo Báo cáo tự đánh giá Chương trình đào tạo Ngành Công nghệ thông tin)")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        run = p.runs[0]
        run.font.size = Pt(14)

        # Khung ghi chú
        table = document.add_table(rows=1, cols=1)
        table.autofit = True
        cell = table.cell(0, 0)
        # Thiết lập nền cho ô (HEX #4F81BD)
        shading_elm = parse_xml(r'<w:shd {} w:fill="4F81BD"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)

        # para = cell.paragraphs[0]
        # para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # run = para.add_run("Lưu ý:\n")
        # run.bold = True
        # run.italic = True

        # Ghi nội dung
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(
            "Lưu ý:\n"
            "- Tài liệu này là tài sản riêng của Khoa HTTT&VT được dùng để làm báo cáo tự đánh giá chương trình đào tạo.\n"
            "- Tài liệu này không được chia sẻ cho bất kỳ cá nhân hoặc tổ chức nào nếu không được sự đồng ý của Khoa HTTT&VT.\n"
            "- Không được thực hiện các hành vi sao chụp, phát tán tài liệu dưới mọi hình thức.\n"
            "- Những cá nhân không liên quan đề nghị không đọc tài liệu này."
        )
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(255, 255, 255)  # Màu chữ trắng

        # Địa điểm và thời gian
        for _ in range(2):
            document.add_paragraph()
        p = document.add_paragraph()
        p = document.add_paragraph(f"TP, Hồ Chí Minh, tháng {int(current_month)}  năm {current_year}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        run = p.runs[0]
        run.font.size = Pt(14)
        
        # Đặt đoạn văn cuối cùng xuống cuối trang
        p.paragraph_format.space_after = Pt(0)  # Loại bỏ khoảng cách sau đoạn
        p.paragraph_format.page_break_before = False  # Đảm bảo không ngắt trang trước đoạn

        document.add_page_break()

        document.add_section(WD_SECTION.NEW_PAGE)
        new_section = document.sections[2]
        # Ngắt liên kết header/footer với section trước
        new_section.header.is_linked_to_previous = False
        new_section.footer.is_linked_to_previous = False
        
        # === Header Section 1 ===
        header = new_section.header
        header_paragraph = header.paragraphs[0]
        header_paragraph.text = "Danh mục minh chứng báo cáo tự đánh giá CTĐT ngành Công nghệ thông tin"
        header_paragraph.paragraph_format.left_indent = Inches(0)
        header_paragraph.runs[0].font.name = 'Times New Roman'
        header_paragraph.runs[0].font.size = Pt(13)
        header_paragraph.runs[0].font.italic = True  # In nghiêng
        
        # Thêm đường kẻ phân chia dưới header
        header0_border = OxmlElement('w:pBdr')
        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(qn('w:val'), 'single')
        bottom_border.set(qn('w:sz'), '12')
        bottom_border.set(qn('w:space'), '4')
        bottom_border.set(qn('w:color'), 'auto')
        header0_border.append(bottom_border)
        header_paragraph._element.get_or_add_pPr().append(header0_border)

        # === Footer Section 1 ===
        footer = new_section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.text = "Khoa Công nghệ thông tin - HCMUNRE"
        footer_paragraph.paragraph_format.left_indent = Inches(0)
        footer_paragraph.runs[0].font.size = Pt(12)
        footer_paragraph.runs[0].font.italic = True  # In nghiêng
        # Thêm đường kẻ phân chia trên footer
        footer0_border = OxmlElement('w:pBdr')
        top_border = OxmlElement('w:top')
        top_border.set(qn('w:val'), 'single')
        top_border.set(qn('w:sz'), '12')
        top_border.set(qn('w:space'), '4')
        top_border.set(qn('w:color'), 'auto')
        footer0_border.append(top_border)
        footer_paragraph._element.get_or_add_pPr().append(footer0_border)
        
        # ===== MỤC LỤC =====
        heading = document.add_paragraph("MỤC LỤC", style='Heading 1')
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.runs[0]
        run.font.size = Pt(18)
        # TOC - Chèn dòng thông báo thay vì dùng XML (để tránh lỗi)
        toc_para = document.add_paragraph("Mục lục sẽ được hiển thị tại đây khi cập nhật thủ công.")
        toc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = toc_para.runs[0]
        run.font.size = Pt(14)
        document.add_page_break()
        
        # Thêm tiêu đề
        title = document.add_paragraph("DANH MỤC MINH CHỨNG")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.runs[0]
        run.bold = True
        run.font.size = Pt(14)

        # # Tạo bảng
        # table = document.add_table(rows=1, cols=5)
        # table.style = 'Table Grid'

        # Header in đậm STT trong tiêu chí
        headers = ['Tiêu chí', 'STT', 'Mã minh chứng', 'Tên minh chứng', 'Số, ngày ban hành, hoặc thời điểm khảo sát, điều tra, phỏng vấn, quan sát,…', 'Nơi ban hành hoặc nhóm, cá nhân thực hiện', 'Ghi chú']
        
        # =========================================================================
        
        # Tạo bảng với số cột tương ứng với headers
        table = document.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        
        # # Căn chỉnh độ rộng các cột (tùy chỉnh theo nhu cầu)
        # table.autofit = False  # Để giữ độ rộng đc set
        # column_widths = [Inches(0.2), Inches(0.2), Inches(0.2), Inches(0.5), Inches(0.5), Inches(1.2), Inches(0.2)]
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            # hdr_cells[i].width = column_widths[i]
            para = hdr_cells[i].paragraphs[0]
            run = para.add_run(header)
            run.bold = True

        current_standard = None
        current_criterion = None
        current_attest_id = None
        stt_within_criterion = 0

        criterion_cells = []
        attest_cells = []
        stt_cells = []

        sorted_queryset = queryset.order_by('criterion__standard', 'criterion_id', 'attest_id')


        # def merge_and_clear(cells):
        #     if len(cells) > 1:
        #         for cell in cells[1:]:
        #             for para in cell.paragraphs:
        #                 para.clear()  # Xóa nội dung
        #         cells[0].merge(cells[-1])  # Merge toàn bộ

        # for obj in sorted_queryset:
        #     if obj.criterion.standard != current_standard:
        #         current_standard = obj.criterion.standard
        #         std_row = table.add_row()
        #         std_cell = std_row.cells[0]
        #         std_cell.merge(std_row.cells[4])
        #         run = std_cell.paragraphs[0].add_run(f"{current_standard.title}")#tiêu chuẩn
        #         run.bold = True

        #     if obj.criterion != current_criterion:
        #         merge_and_clear(criterion_cells)
        #         criterion_cells = []
        #         current_criterion = obj.criterion
        #         stt_within_criterion = 0

        #     if obj.attest_id != current_attest_id:
        #         merge_and_clear(attest_cells)
        #         merge_and_clear(stt_cells)
        #         attest_cells = []
        #         stt_cells = []
        #         current_attest_id = obj.attest_id
        #         stt_within_criterion += 1

        #     row = table.add_row()
        #     cells = row.cells

        #     # Tiêu chí
        #     cells[0].text = str(obj.criterion)
        #     criterion_cells.append(cells[0])
        #     if len(criterion_cells) == 1:
        #         run = cells[0].paragraphs[0].runs[0]
        #         run.bold = True

        #     # STT
        #     cells[1].text = str(stt_within_criterion)
        #     stt_cells.append(cells[1])

        #     # Mã minh chứng
        #     cells[2].text = obj.attest_id
        #     attest_cells.append(cells[2])

        #     # Tên minh chứng
        #     cells[3].text = obj.title or ''
            
        #     # Số ngày ban hành
        #     cells[4].text = obj.body or ''

        #     # Nơi ban hành
        #     cells[5].text = obj.performer or ''
            
        #     # Ghi chú
        #     cells[6].text = obj.note or ''
        # # Merge lần cuối
        # merge_and_clear(criterion_cells)
        # merge_and_clear(attest_cells)
        # merge_and_clear(stt_cells)

        def merge_and_clear(cells):
            if len(cells) > 1:
                for cell in cells[1:]:
                    for para in cell.paragraphs:
                        para.clear()
                cells[0].merge(cells[-1])

        for obj in sorted_queryset:
            # Khi gặp Tiêu chuẩn mới
            if obj.criterion.standard != current_standard:
                # Merge các phần trước nếu có
                merge_and_clear(criterion_cells)
                merge_and_clear(attest_cells)
                merge_and_clear(stt_cells)
                criterion_cells, attest_cells, stt_cells = [], [], []

                current_standard = obj.criterion.standard
                std_row = table.add_row()
                std_cell = std_row.cells[0]
                std_cell.merge(std_row.cells[4])
                run = std_cell.paragraphs[0].add_run(f"{current_standard.title}")
                run.bold = True

                # Reset criterion & attest
                current_criterion = None
                current_attest_id = None

            # Khi gặp Tiêu chí mới
            if obj.criterion != current_criterion:
                merge_and_clear(criterion_cells)
                criterion_cells = []
                current_criterion = obj.criterion
                stt_within_criterion = 0
                current_attest_id = None  # Reset để phân nhóm minh chứng mới

            # Khi gặp mã minh chứng mới (trong cùng tiêu chí)
            if obj.attest_id != current_attest_id:
                merge_and_clear(attest_cells)
                merge_and_clear(stt_cells)
                attest_cells = []
                stt_cells = []
                current_attest_id = obj.attest_id
                stt_within_criterion += 1

            # Tạo hàng mới
            row = table.add_row()
            cells = row.cells

            # Cột Tiêu chí
            cells[0].text = str(obj.criterion)
            criterion_cells.append(cells[0])
            if len(criterion_cells) == 1:
                run = cells[0].paragraphs[0].runs[0]
                run.bold = True

            # Cột STT
            cells[1].text = str(stt_within_criterion)
            stt_cells.append(cells[1])

            # Cột Mã minh chứng
            cells[2].text = obj.attest_id
            attest_cells.append(cells[2])

            # Các cột còn lại
            cells[3].text = obj.title or ''
            cells[4].text = obj.body or ''
            cells[5].text = obj.performer or ''
            cells[6].text = obj.note or ''

        # Merge cuối cùng
        merge_and_clear(criterion_cells)
        merge_and_clear(attest_cells)
        merge_and_clear(stt_cells)

        # # Tạo bảng với số cột tương ứng với headers
        # headers = ['Tiêu chí', 'STT', 'Mã minh chứng', 'Tên minh chứng', 'Số, ngày ban hành, hoặc thời điểm khảo sát, điều tra, phỏng vấn, quan sát,…', 'Nơi ban hành hoặc nhóm, cá nhân thực hiện', 'Ghi chú']
        # table = document.add_table(rows=1, cols=len(headers))
        # table.style = 'Table Grid'

        # # Thiết lập tiêu đề bảng
        # hdr_cells = table.rows[0].cells
        # for i, header in enumerate(headers):
        #     para = hdr_cells[i].paragraphs[0]
        #     run = para.add_run(header)
        #     run.bold = True

        # # Khởi tạo biến để theo dõi tiêu chuẩn, tiêu chí, và mã minh chứng hiện tại
        # current_standard = None
        # current_criterion = None
        # current_attest_id = None
        # stt_within_criterion = 0

        # # Sắp xếp queryset
        # sorted_queryset = queryset.order_by('criterion__standard', 'criterion_id', 'attest_id')

        # def merge_and_clear(cells):
        #     if len(cells) > 1:
        #         for cell in cells[1:]:
        #             for para in cell.paragraphs:
        #                 para.clear()  # Xóa nội dung
        #         cells[0].merge(cells[-1])  # Merge toàn bộ
        
        # # Duyệt qua từng đối tượng trong queryset để điền dữ liệu vào bảng
        # for obj in sorted_queryset:
        #     # Thêm hàng tiêu chuẩn nếu tiêu chuẩn thay đổi
        #     if obj.criterion.standard != current_standard:
        #         current_standard = obj.criterion.standard
        #         std_row = table.add_row()
        #         std_cell = std_row.cells[0]
        #         # Gộp các ô trong hàng tiêu chuẩn
        #         std_cell.merge(std_row.cells[len(headers) - 1])
        #         run = std_cell.paragraphs[0].add_run(f"{current_standard.title}")
        #         run.bold = True

        #     # Kiểm tra nếu tiêu chí thay đổi
        #     if obj.criterion != current_criterion:
        #         current_criterion = obj.criterion
        #         stt_within_criterion = 0

        #     # Kiểm tra nếu mã minh chứng thay đổi
        #     if obj.attest_id != current_attest_id:
        #         current_attest_id = obj.attest_id
        #         stt_within_criterion += 1

        #     # Thêm hàng mới vào bảng
        #     row = table.add_row()
        #     cells = row.cells

        #     # Tiêu chí: Chỉ hiển thị nếu là lần đầu tiên xuất hiện
        #     if obj.criterion == current_criterion and row._index > 1:  # Kiểm tra nếu không phải hàng đầu tiên
        #         displayed_criterion = ""
        #         for previous_row in table.rows[:row._index]:
        #             if previous_row.cells[0].text == str(obj.criterion):
        #                 displayed_criterion = ""
        #                 break
        #             else:
        #                 displayed_criterion = str(obj.criterion)
        #         cells[0].text = displayed_criterion
        #     else:
        #         cells[0].text = str(obj.criterion)
        #     run = cells[0].paragraphs[0].runs[0]
        #     run.bold = True

        #     # STT: Hiển thị số thứ tự trong tiêu chí
        #     if obj.attest_id == current_attest_id and row._index > 1:
        #         displayed_stt = ""
        #         for previous_row in table.rows[:row._index]:
        #             if previous_row.cells[2].text == obj.attest_id:
        #                 displayed_stt = ""
        #                 break
        #             else:
        #                 displayed_stt = str(stt_within_criterion)
        #         cells[1].text = displayed_stt
        #     else:
        #         cells[1].text = str(stt_within_criterion)

        #     # Mã minh chứng: Chỉ hiển thị nếu là lần đầu tiên
        #     if obj.attest_id == current_attest_id and row._index > 1:
        #         displayed_attest = ""
        #         for previous_row in table.rows[:row._index]:
        #             if previous_row.cells[2].text == obj.attest_id:
        #                 displayed_attest = ""
        #                 break
        #             else:
        #                 displayed_attest = obj.attest_id
        #         cells[2].text = displayed_attest
        #     else:
        #         cells[2].text = obj.attest_id

        #     # Tên minh chứng
        #     cells[3].text = obj.title or ''

        #     # Số ngày ban hành
        #     cells[4].text = obj.body or ''

        #     # Nơi ban hành
        #     cells[5].text = obj.performer or ''

        #     # Ghi chú
        #     cells[6].text = obj.note or ''
        
        # ======================================================================

        # # ===== DANH MỤC MINH CHỨNG DÙNG CHUNG =====
        # document.add_page_break()
        # title = document.add_paragraph("DANH MỤC CÁC VĂN BẢN MINH CHỨNG DÙNG CHUNG")
        # title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # run = title.runs[0]
        # run.bold = True
        # run.font.size = Pt(14)

        # headers_common = ['STT', 'Tên văn bản', 'Mã minh chứng']
        # table_common = document.add_table(rows=1, cols=3)
        # table_common.style = 'Table Grid'
        # hdr_cells = table_common.rows[0].cells
        # for i, header in enumerate(headers_common):
        #     para = hdr_cells[i].paragraphs[0]
        #     run = para.add_run(header)
        #     run.bold = True

        # from .models import common_attest
        # common_queryset = common_attest.objects.order_by('common_attest_stt')

        # for cm_attest in common_queryset:
        #     row = table_common.add_row()
        #     cells = row.cells
        #     cells[0].text = cm_attest.common_attest_stt
        #     cells[1].text = cm_attest.title
        #     cells[2].text = cm_attest.common_attest_id


        # Xuất file
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="DanhMucMinhChung.docx"'
        document.save(response)
        return response
    
    
    

@admin.register(common_attest)
# class common_attestAdmin(ImportExportActionModelAdmin, admin.ModelAdmin):
class common_attestAdmin(admin.ModelAdmin):
    form = CommonAttestForm
    inlines = [PhotoCommonAttestInline]
    
    def save_related(self, request, form, formsets, change):
        super().save_related(request, form, formsets, change)
        form.save_photos(form.instance)
    
    list_display = ('common_attest_id_name','common_attest_stt', 'title', 'body', 'performer')
    list_display_links = ('title',)
    # list_filter = (
    #     #'criterion',
    #     # for ordinary fields
    #     #('criterion', DropdownFilter),
    #     # for choice fields
    #     #('criterion', ChoiceDropdownFilter),
    #     # for related fields
    #     ('criterion', RelatedDropdownFilter),
    # )
    
    ordering = ( 'common_attest_id','common_attest_stt')
    search_fields = ('title', 'performer')
    # prepopulated_fields = {
    #     'slug': ['common_attest_id','common_attest_stt'],
    #     # 'common_attest_id':['box','criterion'],
    # }
    
    class Media:
        js = ('../static/js/custom_admin.js', 'https://cdnjs.cloudflare.com/ajax/libs/speakingurl/14.0.1/speakingurl.min.js')  # Đường dẫn file JS
        css = {
            'all': ('../static/css/custom_admin.css',)
        }
    def save_model(self, request, obj, form, change):
        """
        Ghi đè save_model cập nhật tất cả các attest liên kết với common_attest.
        """
        super().save_model(request, obj, form, change)  # 🔹 Đảm bảo obj đã được lưu trước khi lấy pk
        if change :
            action_type = "Cập nhật minh chứng dùng chung"
        else :
            action_type = "Thêm mới minh chứng dùng chung"
        admin_url = request.build_absolute_uri(reverse('admin:CTDT_common_attest_change', args=[obj.pk]))
        
        # Tìm tất cả các attest liên quan tới common_attest hiện tại
        related_attests = attest.objects.filter(common_attest=obj)
        # Cập nhật các trường trong attest liên quan nếu cần
        for attest_instance in related_attests:
            attest_instance.title = obj.title  # Đồng bộ trường `title`
            attest_instance.body = obj.body  
            attest_instance.performer = obj.performer  
            attest_instance.note = obj.note  
            attest_instance.slug = obj.slug  
            # attest_instance.image = obj.image  
            
            attest_instance.criterion = obj.criterion  
            attest_instance.box = obj.box  
            attest_instance.save()  # Lưu thay đổi cho từng instance
            
        
        # def update_photos():
        #     related_attests1 = attest.objects.filter(common_attest=obj)

        #     for attest_instance1 in related_attests1:
        #         attest_instance1.photos.all().delete()  # Xóa ảnh cũ
                    
        #         for photo_common1 in obj.photos.all():
        #             PhotoAttest.objects.create(show=attest_instance1, photo=photo_common1.photo)    
        # transaction.on_commit(update_photos(obj))  # Đảm bảo chạy sau khi commit database
        
        transaction.on_commit(lambda: ActionConvert.update_photos(obj))  # Đảm bảo chạy sau khi commit database
        
        # Xử lý các PhotoAttest instances
        if 'photocommonattest_set' in form.cleaned_data:  # Kiểm tra xem có PhotoAttest instances trong form không
            for photo in form.cleaned_data['photocommonattest_set']:
                try:
                    self.inlines[0].clean_photo(photo)  # Gọi hàm clean_photo nếu cần
                except forms.ValidationError as e:
                    raise forms.ValidationError(e)  # Ném lỗi nếu có vấn đề
                photo.save()  # Lưu PhotoAttest instance
        
        # # EmailNotification.send_common_attest_email(request, [obj], action_type, admin_url)
        # transaction.on_commit(lambda: 
        #     EmailNotification.send_common_attest_email(request, [obj], action_type, admin_url)
        # )
    @admin.display(description="Mã minh chứng")
    
    def common_attest_id_name(self, obj):
        if self.model.objects.filter(common_attest_id=obj.common_attest_id, pk__lt=obj.pk).exists():
            return ""
        return f"{obj.common_attest_id}".upper()

    def delete_model(self, request, obj):
        """ Gửi email khi xóa """
        # EmailNotification.send_common_attest_email(request, [obj], "Xóa minh chứng dùng chung", "Delete")
        
        ActionConvert.delete_attests(obj)  # xóa ảnh
        super().delete_model(request, obj)
    
    def delete_queryset(self, request, queryset):
        """ Gửi email chứa danh sách attest bị xóa trước khi xóa """
        
        # delete photo
        for common_attest in queryset:
            ActionConvert.delete_attests(common_attest)  # Xóa ảnh liên quan
            for photo_attest in common_attest.photos.all():  # Lấy ảnh liên kết
                if photo_attest.photo:
                    try:
                        thumbnail_path = get_thumbnailer(photo_attest.photo)['small'].path
                        if os.path.isfile(thumbnail_path):
                            os.remove(thumbnail_path)
                    except Exception:
                        pass  # Bỏ qua nếu không có thumbnail

                    if os.path.isfile(photo_attest.photo.path):
                        os.remove(photo_attest.photo.path)
                    
                    # Kiểm tra thư mục tồn tại trước khi xóa
                    folder1 = os.path.dirname(photo_attest.photo.path)
                    # xóa index
                    remove_image_from_index(photo_attest.photo.path)
                    if os.path.exists(folder1) and not os.listdir(folder1):  # Kiểm tra thư mục rỗng
                        shutil.rmtree(folder1)
                    # # Kiểm tra và xóa folder nếu rỗng (loại bỏ file ẩn nếu cần)
                    # remaining_files = [f for f in os.listdir(folder1) if not f.startswith('.')]
                    # if not remaining_files:
                    #     shutil.rmtree(folder1)
                    
                    
        
        # EmailNotification.send_common_attest_email(request, queryset, "Xóa minh chứng dùng chung", "Delete")
        
        # Gọi phương thức mặc định để xóa các attest
        super().delete_queryset(request, queryset)
        

# chưa cập nhật được slug , ẩn trường slug khi chỉnh sửa, thêm mới, cập nhật trường trong tiêu chí ==> xong

# Không gửi mail cho ad khi thao tác với database, cần tạo mới 1 table lưu các thông tin chỉnh sửa của người dùng 
# đến khi nào người dùng gửi thông báo cho admin thì sẽ lấy hết thông tin trong bảng đó gửi cho admin
