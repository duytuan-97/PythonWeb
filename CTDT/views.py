from django.shortcuts import render

# Create your views here.
from django.http import HttpResponse, HttpResponseRedirect

from CTDT.models import Post
from CTDT.models import box, standard, criterion, attest, common_attest

from docx import Document

from django.utils.text import slugify

import re


from .notifications import EmailNotification


# Create your views here.
def index(request):
#    response = HttpResponse()
#    response.writelines('<h1>Xin chào</h1>')
#    response.write('Đây là app CTDT')
#    return response
    # return render(request, 'pages/home.html')
    
    # Xu Ly với Database
    # a = Post()
    # a.title = 'First Title'
    # a.body = 'Hello World'
    # a.save()

    # return render(request, 'view/index.html')
    return render(request, 'Pages/dashboard.html')

def dashboard_test(request):

    return render(request, 'Pages/dashboard_test.html')

def posts_list_test(request):
    posts = Post.objects.all().order_by('-date')
    return render(request, 'Pages/posts_list_test.html', {'posts': posts})

def post_page(request, slug):
    post = Post.objects.get(slug=slug)
    return render(request, 'Pages/post_page_test.html', {'post': post})


# xác thực mới vào trang được
from django.contrib.auth.decorators import login_required

@login_required(login_url="/users/login/")#kiểm tra xem người dùng đã đăng nhập chưa, nếu chưa quay lại trang đăng nhập
def posts_new_test(request):
    # return render(request, 'Pages/new_post.html')
    return render(request, 'Pages/posts_new_test.html')
    

#Xử lý File
from .forms import UploadFileForm
from .tools import process_file
from django.contrib import messages
from django.shortcuts import redirect

from .forms import FileUploadForm
import os

# def upload_file(request):
#     if request.method == 'POST':
#         form = UploadFileForm(request.POST, request.FILES)
#         if form.is_valid():
#             file = request.FILES['file']  

#             # Lưu file vào thư mục
#             with open('media/' + file.name, 'wb+') as destination:
#                 for chunk in file.chunks():
#                     destination.write(chunk)

#             # Gọi hàm xử lý file
#             process_file('media/' + file.name)

#             messages.success(request, 'Hành động đã được thực hiện thành công!')
#             return redirect('dashboard_test')
#             # return HttpResponseRedirect('/success/')  # Chuyển hướng đến trang thông báo thành công
#     else:
#         form = UploadFileForm()
#     return render(request, 'upload.html', {'form': form})

def upload_file(request):
    if request.method == 'POST':
        form = FileUploadForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES['file']
            file_path = os.path.join('media', uploaded_file.name)
            with open(file_path, 'wb+') as destination:
                for chunk in uploaded_file.chunks():
                    destination.write(chunk)
            result_message = process_file(file_path)
            # # return HttpResponse(result_message)
            messages.success(request, 'Đã tạo thư mục theo cấu trúc thành công!')
            return redirect('posts:upload_file')
    else:
        form = FileUploadForm()

    # return render(request, 'Pages/upload.html', {'form': form})
    return render(request, 'admin/test/test.html', {'form': form})

def convert_code(text):
    parts = text.split(".")
    return ".".join(str(int(part)) if part.isdigit() else part for part in parts)

def import_word(request):
    if request.method == 'POST':
        list_attest = []
        word_file = request.FILES['word_file']
        try:
            document = Document(word_file)
            tables = document.tables # Lấy tất cả các bảng trong file

            if tables: # Kiểm tra xem có bảng nào không
                table = tables[0] # Lấy bảng đầu tiên (nếu có nhiều bảng)
                stt = 1
                
                # Bỏ qua hàng tiêu đề (hàng đầu tiên)
                for row in table.rows[1:]: # bắt đầu từ hàng thứ 2
                    cells = row.cells
                    try:
                        
                        tieu_chi = cells[0].text.strip() # Cột "Mã tiêu chí" (index 0)
                        ma_minh_chung = cells[2].text.strip() # Cột "Mã minh chứng" (index 2)
                        title = cells[3].text.strip() # Cột "Tên minh chứng" (index 3)
                        # ... Lấy dữ liệu từ các cột khác tương tự
                        so_ngay_ban_hanh = cells[4].text.strip()
                        noi_ban_hanh = cells[5].text.strip()
                        ghi_chu = cells[6].text.strip()
                        is_common1 = False
                        

                        if not ma_minh_chung: # kiểm tra nếu mã minh chứng bị trống thì bỏ qua hàng này
                            continue
                        if cells[2].text.strip() == cells[3].text.strip():
                            messages.warning(request, f"{ma_minh_chung}")
                            continue
                        box_id_str, standard_str, criterion_str, attest_str = ma_minh_chung.split('.')
                        box_id = int(box_id_str[1:])

                        match = re.search(r'\d+\.\d+', tieu_chi)  # Tìm mẫu số chấm số
                        # ma_tieu_chi = match.group() if match else None
                        ma_tieu_chi = match.group() if match else None
                        # ma_tieu_chi = convert_code(ma_tieu_chi)
                        

                        # Kiểm tra Box
                        try:
                            box1 = box.objects.get(pk=box_id)
                        except box1.DoesNotExist:
                            messages.error(request, f"Box H{box_id} không tồn tại.")
                            continue

                        # Kiểm tra Standard (cần xử lý thêm để lấy Standard code từ Tiêu chí)
                        try:
                            tieu_chi_text = cells[0].text.strip()
                            standard_str = tieu_chi_text.split('.')[0].split(' ')[-1] # lấy số tiêu chuẩn từ cột tiêu chí
                            # standard = standard.objects.get(code=standard_str)
                            standard1 = standard.objects.get(id=int(standard_str))
                            
                        except standard.DoesNotExist:
                            messages.error(request, f"Tiêu chuẩn {standard_str} không tồn tại.")
                            continue
                        except:
                            messages.error(request, f"Lỗi xử lý cột tiêu chuẩn")
                            continue

                        # Kiểm tra Criterion
                        try:
                            criterion1 = criterion.objects.get(pk=ma_tieu_chi)
                        except criterion1.DoesNotExist:
                            messages.error(request, f"Tiêu chí {ma_tieu_chi} không tồn tại.")
                            continue
                        
                        # Kiểm tra Minh chứng (Attest)
                        ma_minh_chung = ma_minh_chung.strip()
                        so_ngay_ban_hanh = so_ngay_ban_hanh.strip()
                        if attest.objects.filter(attest_id=ma_minh_chung, body=so_ngay_ban_hanh).exists():
                            messages.warning(request, f"Minh chứng {ma_minh_chung} đã tồn tại.")
                            continue
                        if attest.objects.filter(attest_id=ma_minh_chung).exists():
                            stt += 1
                        else:
                            stt = 1

                        
                        # Kiểm tra minh chứng dùng chung
                        match = re.search(r'\bDC\b', ghi_chu)
                        common_evidence = None
                        if not match:
                            print("Ghi chú không chứa ký tự 'DC'.")
                        else:
                            print("Ghi chú có chứa ký tự 'DC'.")
                            try:
                                # Truy vấn bản ghi dùng chung
                                # convert_code(ma_minh_chung)
                                common_evidence = common_attest.objects.get(common_attest_id=ma_minh_chung, body=so_ngay_ban_hanh)
                                # common_attest_instance = common_attest.objects.get(common_attest_id=ma_minh_chung)
                                print("Tìm thấy minh chứng dùng chung:", common_evidence)
                            except common_attest.DoesNotExist:
                                messages.error(request, f"Không tìm thấy minh chứng dùng chung với mã và số thứ tự đã cho.")
                                return render(request, 'admin/import_word.html')
                            
                            # So sánh các trường chung
                            minh_chung_check = {
                                "common_attest_id": ma_minh_chung,
                                # "common_attest_stt": stt,
                                "title": title,
                                "body": so_ngay_ban_hanh,
                                "performer": noi_ban_hanh,
                                # "note": ghi_chu,
                            }
                            differences = {
                                field: (getattr(common_evidence, field), minh_chung_check[field])
                                for field in minh_chung_check.keys()
                                if hasattr(common_evidence, field) and getattr(common_evidence, field) != minh_chung_check[field]
                            }

                            if differences:
                                messages.error(request, "Có sự khác biệt trong các trường sau:")
                                for field, (common_value, word_value) in differences.items():
                                    messages.error(f"{field}: Dùng chung = '{common_value}', Từ Word = '{word_value}'")
                                return render(request, 'admin/import_word.html')
                            is_common1 = True
                        #Lỗi trùng dữ liệu với code : common_attest = common_evidence
                        # Tạo và lưu Minh chứng mới
                        attest1 = attest(
                            attest_id=ma_minh_chung,
                            title=title,
                            box=box1,
                            criterion=criterion1,
                            body=so_ngay_ban_hanh,
                            performer=noi_ban_hanh,
                            note = ghi_chu,
                            attest_stt = stt,
                            slug=slugify(ma_minh_chung)+ "_" + str(stt),  # Tạo slug từ attest_id
                            # slug=slugify(ma_minh_chung),  # Tạo slug từ attest_id
                            # photo='fallback.jpeg',  # Sử dụng hình mặc định nếu không có hình được cung cấp
                            is_common = is_common1,
                            common_attest = common_evidence
                        )
                        attest1.save()
                        print(attest1)
                        list_attest.append(attest1)
                        print("list_attest:   ")
                        print(list_attest)
                        messages.success(request, f"Đã import thành công minh chứng {ma_minh_chung}.")

                    except ValueError as ve:
                        messages.error(request, f"Lỗi định dạng dữ liệu trong bảng: {ve}")
                    except IndexError as ie:
                        messages.error(request, f"Lỗi thiếu cột trong bảng: {ie}. Kiểm tra cấu trúc bảng")
                    except Exception as e:
                        messages.error(request, f"Lỗi không xác định khi import minh chứng: {e}")

                
                EmailNotification.send_attest_email(request, list_attest, "Thêm mới minh chứng")
            else:
                messages.error(request, "Không tìm thấy bảng nào trong file DOCX.")

        except Exception as e:
            messages.error(request, f"error: {e}")
    # else:
    #     form = FileUploadForm()
    return render(request, 'admin/import_word.html')

# def import_word(request):
#     if request.method == 'POST':
#         docx_file = request.FILES['docx_file']
#         try:
#             document = Document(docx_file)
#             tables = document.tables # Lấy tất cả các bảng trong file

#             if tables: # Kiểm tra xem có bảng nào không
#                 table = tables[0] # Lấy bảng đầu tiên (nếu có nhiều bảng)
#                 # Bỏ qua hàng tiêu đề (hàng đầu tiên)
#                 for row in table.rows[1:]: # bắt đầu từ hàng thứ 2
#                     cells = row.cells
#                     try:
#                         ma_minh_chung = cells[2].text.strip() # Cột "Mã minh chứng" (index 2)
#                         title = cells[3].text.strip() # Cột "Tên minh chứng" (index 3)
#                         # ... Lấy dữ liệu từ các cột khác tương tự
#                         so_ngay_ban_hanh = cells[4].text.strip()
#                         noi_ban_hanh = cells[5].text.strip()
#                         ghi_chu = cells[6].text.strip()

#                         if not ma_minh_chung: # kiểm tra nếu mã minh chứng bị trống thì bỏ qua hàng này
#                             continue

#                         box_id_str, standard_str, criterion_str, attest_str = ma_minh_chung.split('.')
#                         box_id = int(box_id_str[1:])

#                         # Kiểm tra Box
#                         try:
#                             box = box.objects.get(pk=box_id)
#                         except box.DoesNotExist:
#                             messages.error(request, f"Box H{box_id} không tồn tại.")
#                             continue

#                         # Kiểm tra Standard (cần xử lý thêm để lấy Standard code từ Tiêu chí)
#                         try:
#                             tieu_chi_text = cells[0].text.strip()
#                             standard_str = tieu_chi_text.split('.')[0].split(' ')[-1] # lấy số tiêu chuẩn từ cột tiêu chí
#                             standard = standard.objects.get(code=standard_str)
#                         except standard.DoesNotExist:
#                             messages.error(request, f"Tiêu chuẩn {standard_str} không tồn tại.")
#                             continue
#                         except:
#                             messages.error(request, f"Lỗi xử lý cột tiêu chí")
#                             continue

#                         # Kiểm tra Criterion
#                         try:
#                             criterion = criterion.objects.get(code=criterion_str)
#                         except criterion.DoesNotExist:
#                             messages.error(request, f"Tiêu chí {criterion_str} không tồn tại.")
#                             continue
                        
#                         # Kiểm tra Minh chứng (Attest)
#                         if attest.objects.filter(attest_id=ma_minh_chung).exists():
#                             messages.warning(request, f"Minh chứng {ma_minh_chung} đã tồn tại.")
#                             continue

#                         # Tạo và lưu Minh chứng mới
#                         attest = attest(
#                             attest_id=ma_minh_chung,
#                             title=title,
#                             box=box,
#                             criterion=criterion,
#                             attest_stt=attest_str,
#                             body=so_ngay_ban_hanh,
#                             performer=noi_ban_hanh,
#                             note = ghi_chu
#                             # ... Các trường khác
#                         )
#                         attest.save()
#                         messages.success(request, f"Đã import thành công minh chứng {ma_minh_chung}.")

#                     except ValueError as ve:
#                         messages.error(request, f"Lỗi định dạng dữ liệu trong bảng: {ve}")
#                     except IndexError as ie:
#                         messages.error(request, f"Lỗi thiếu cột trong bảng: {ie}. Kiểm tra cấu trúc bảng")
#                     except Exception as e:
#                         messages.error(request, f"Lỗi không xác định khi import minh chứng: {e}")

#             else:
#                 messages.error(request, "Không tìm thấy bảng nào trong file DOCX.")

#         except Exception as e:
#             messages.error(request, f"Lỗi khi đọc file DOCX: {e}")

#     return render(request, 'your_template.html')

from django.http import JsonResponse
from django.shortcuts import get_object_or_404

def get_common_attest_data(request, pk):
    common = get_object_or_404(common_attest, pk=pk)
    # hhhh = common.slug
    # h11 = hhhh.id
    return JsonResponse({
        "common_attest_id": common.common_attest_id,
        "common_attest_stt": common.common_attest_stt,
        "performer": common.performer,
        "slug": common.slug,
        # "image": common.image.name,
        "criterion": common.criterion.id,
        "box": common.box.id,
        "title": common.title,
        "body": common.body,
        # Thêm các trường khác nếu cần
    })
    
def custom_admin_view(request):
    return redirect('/admin/')  # Chuyển hướng đến trang Admin
